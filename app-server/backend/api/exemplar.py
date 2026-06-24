"""
范文解构 API：上传 .docx 范文 → 解析标题+正文内容 → 持久化存储。
WHY: 与 template.py（纯骨架抽取，content 为空）不同，
     exemplar 保留各章节的完整正文内容，作为双路 Prompt 的风格参考源。
"""
from __future__ import annotations

import json
import re
import uuid
import tempfile
import os
import logging
from pathlib import Path
from typing import List

from fastapi import APIRouter, File, UploadFile, HTTPException, Depends
from pydantic import BaseModel
from docx import Document
from lxml import etree

from core.auth_deps import get_current_user
from core.config import settings
from core.heading_utils import determine_level_from_text
from core.project_access import require_project_access

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/exemplar", tags=["范文解构"])

# WHY: 范文数据独立于模板（template），存放在 data/exemplars/ 下
EXEMPLARS_DIR = Path(settings.DATA_DIR) / "exemplars"
EXEMPLARS_DIR.mkdir(parents=True, exist_ok=True)


def _table_to_markdown(tbl_element, doc) -> str:
    """
    将 Word 表格 XML 元素转为 Markdown 格式。
    WHY: 范文中的表格也是重要的风格参考，
         AI 需要看到范文的表格排版才能模仿生成类似结构。
    """
    from docx.table import Table
    try:
        table = Table(tbl_element, doc)
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")
        if rows:
            header = rows[0]
            separator = "| " + " | ".join(
                ["---"] * len(table.rows[0].cells)
            ) + " |"
            md = header + "\n" + separator
            if len(rows) > 1:
                md += "\n" + "\n".join(rows[1:])
            return md
    except Exception as e:
        logger.warning(f"表格转 Markdown 失败: {e}")
    return ""


def _has_image(element) -> bool:
    """
    检测段落元素中是否包含嵌入图片（通过 blipFill XML 节点判断）。
    WHY: 范文中的图件位置需要记录为占位符，
         指导 AI 在对应位置输出 [插入图件：...] 标记。
    """
    try:
        xml_str = etree.tostring(element, encoding="unicode")
        return "blipFill" in xml_str
    except Exception:
        return False


def _detect_heading_level(para_element, doc) -> int:
    """
    从段落 XML 元素检测标题层级。
    WHY: 正则编号优先于 Word 样式，因为很多文档样式不规范（全用 Heading 1）。
    返回 0 表示非标题（普通正文）。
    """
    from docx.text.paragraph import Paragraph
    para = Paragraph(para_element, doc)
    text = para.text.strip()
    if not text:
        return 0

    # 1. 用文本编号推断层级（最可靠）
    level = determine_level_from_text(text)

    # 2. 编号无法识别时，回退到 Word 样式
    if level == 0:
        style_name = para.style.name
        if style_name.startswith("Heading"):
            try:
                level = int(style_name.replace("Heading", "").strip())
            except ValueError:
                level = 1
        elif "标题" in style_name:
            try:
                level = int(re.search(r"\d+", style_name).group())
            except AttributeError:
                level = 1

    return level if 1 <= level <= 4 else 0


@router.post("/parse")
async def parse_exemplar(
    file: UploadFile = File(...),
    user: dict = Depends(get_current_user),
):
    """
    解析 .docx 范文，提取标题及其正文内容。
    WHY: 使用统一 DOM 遍历（doc.element.body 的所有子元素），
         确保段落和表格按文档中的真实顺序挂到正确的章节下。
         旧方案分别遍历 doc.paragraphs 和 doc.tables 导致表格位置错乱。
    """
    if file.content_type not in [
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ]:
        raise HTTPException(
            status_code=400,
            detail="必须上传 .docx 类型的 Word 文档范文",
        )

    temp_path = os.path.join(
        tempfile.gettempdir(), f"exemplar_{uuid.uuid4().hex}.docx"
    )
    try:
        with open(temp_path, "wb") as f:
            f.write(await file.read())

        doc = Document(temp_path)
        # WHY: 引入 qn 用于判断 XML 元素类型（段落 w:p / 表格 w:tbl）
        from docx.oxml.ns import qn
        from docx.text.paragraph import Paragraph

        sections: list[dict] = []
        current_content_lines: list[str] = []
        image_counter = 0
        table_counter = 0

        # WHY: 统一遍历 body 的所有子元素，段落和表格按文档内的真实顺序处理
        seen_titles = set()  # WHY: 去重，防止 TOC 和正文产生重复章节
        for child in doc.element.body:
            if child.tag == qn('w:p'):
                # ── 段落（可能是标题或正文） ──
                para = Paragraph(child, doc)
                text = para.text.strip()

                # WHY: 过滤 TOC 目录段落 — 与 template.py 保持一致
                style_name = para.style.name
                if style_name.startswith("TOC") or style_name.startswith("toc"):
                    continue

                # 检测图片
                # WHY: 图片占位标记必须与章节正文分离存储。
                #      旧方案将 [图件：图N] 混入 content，导致 LLM 在多个章节中
                #      照搬相同的图片标记（因为 fallback 匹配会轮换到含图片的范文段）。
                #      新方案存入独立的 images 列表，由生成阶段按需精确注入。
                if _has_image(child):
                    image_counter += 1
                    logger.info(f"检测到图片，当前计数: {image_counter}")
                    if sections:
                        sections[-1].setdefault("images", []).append({
                            "index": image_counter,
                            "position": "after_paragraph",
                            "context_hint": current_content_lines[-1][:50] if current_content_lines else "",
                        })
                    continue

                if not text:
                    continue

                level = _detect_heading_level(child, doc)

                if level > 0:
                    # WHY: 清洗标题 — 去掉制表符+页码残留（如 "1.3.1 建设范围\t6"）
                    clean_title = re.sub(r'\t+\d*$', '', text).strip()
                    
                    # WHY: 去重 — 如果清洗后的标题已出现过，跳过
                    dedup_key = re.sub(r'\s+', '', clean_title)
                    if dedup_key in seen_titles:
                        continue
                    seen_titles.add(dedup_key)
                    
                    # 遇到新标题 → 先将之前收集的正文写入上一个 section
                    logger.info(f"检测到新标题 [L{level}]: {clean_title}")
                    if sections and current_content_lines:
                        sections[-1]["content"] = "\n".join(
                            current_content_lines
                        )
                    current_content_lines = []

                    sections.append({
                        "id": str(uuid.uuid4()),
                        "title": clean_title,
                        "level": level,
                        "content": "",
                    })
                else:
                    # 普通正文 → 追加到当前收集器
                    if sections:
                        current_content_lines.append(text)
                    else:
                        logger.debug(f"跳过无归属的正文: {text[:20]}...")

            elif child.tag == qn('w:tbl'):
                # ── 表格 → 立即转为 Markdown 追加到当前 section ──
                table_counter += 1
                logger.debug(f"检测到表格，当前计数: {table_counter}")
                md = _table_to_markdown(child, doc)
                if md and sections:
                    current_content_lines.append(md)
                elif not sections:
                    logger.debug(f"跳过无归属的表格 (计数: {table_counter})")

        # 最后一个 section 的正文收尾
        if sections and current_content_lines:
            sections[-1]["content"] = "\n".join(current_content_lines)

        logger.info(
            f"范文解构完成: {file.filename}, "
            f"共 {len(sections)} 个章节, "
            f"{table_counter} 个表格, "
            f"{image_counter} 处图件"
        )

        return {"filename": file.filename, "sections": sections}

    except Exception as e:
        raise HTTPException(
            status_code=500, detail=f"范文解构失败: {str(e)}"
        )
    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)


# ---------- CRUD: 项目级范文持久化 ----------


class ExemplarData(BaseModel):
    title: str
    sections: list


@router.get("/project/{project_id}")
async def get_exemplar(
    project_id: str, user: dict = Depends(get_current_user)
):
    """读取指定项目绑定的写作范文。
    WHY: 对超长章节截断后返回，防止前端范文预览面板 OOM。
    截断仅影响前端展示，不修改磁盘文件，不影响 AI 生成。
    """
    require_project_access(project_id, user, write=False)
    fp = EXEMPLARS_DIR / f"{project_id}.json"
    if not fp.exists():
        return {"title": "", "sections": []}
    try:
        data = json.loads(fp.read_text(encoding="utf-8"))
        _MAX = 8000
        for s in data.get("sections", []):
            if isinstance(s, dict) and len(s.get("content", "")) > _MAX:
                orig_len = len(s["content"])
                s["content"] = (
                    s["content"][:_MAX]
                    + f"\n\n> ⚠️ 本章节原文共 {orig_len:,} 字，"
                    f"已截断显示前 {_MAX} 字。完整内容已保存。"
                )
        return data
    except Exception:
        return {"title": "", "sections": []}


@router.post("/project/{project_id}")
async def save_exemplar(
    project_id: str,
    data: ExemplarData,
    user: dict = Depends(get_current_user),
):
    """将解构后的范文持久化到项目专属路径。"""
    require_project_access(project_id, user, write=True)
    EXEMPLARS_DIR.mkdir(parents=True, exist_ok=True)
    fp = EXEMPLARS_DIR / f"{project_id}.json"
    fp.write_text(
        json.dumps(data.dict(), ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    logger.info(
        f"范文已保存: project={project_id}, "
        f"title={data.title}, "
        f"sections={len(data.sections)}"
    )

    # WHY: 范文变更后清理旧缓存：Slot 缓存 + 全文预计算草稿缓存
    from core.slot_cache import invalidate
    invalidate(project_id)
    from core.precompute import invalidate_draft_cache
    invalidate_draft_cache(project_id)

    return data.dict()


# ---------- 预计算触发 ----------


class PrecomputeRequest(BaseModel):
    mode: str = "replace"  # generate / replace / clone


@router.post("/project/{project_id}/precompute")
async def trigger_precompute(
    project_id: str,
    body: PrecomputeRequest = PrecomputeRequest(),
    user: dict = Depends(get_current_user),
):
    """
    触发指定模式的预计算。
    WHY: 三种模式对应三种生成策略，各自独立缓存。
    """
    require_project_access(project_id, user, write=True)
    mode = body.mode
    if mode not in ("generate", "replace", "clone"):
        raise HTTPException(400, f"无效模式: {mode}")

    # generate 模式不需要范文，但需要模板
    if mode == "generate":
        template_path = Path(settings.DATA_DIR) / "templates" / f"{project_id}.json"
        if not template_path.exists():
            raise HTTPException(400, "请先创建大纲模板")
    else:
        exemplar_path = EXEMPLARS_DIR / f"{project_id}.json"
        if not exemplar_path.exists():
            raise HTTPException(400, "请先挂载写作范文")

    from core.precompute import invalidate_draft_cache, schedule_precompute
    # 清除该模式的旧缓存后重新计算
    invalidate_draft_cache(project_id, mode)
    schedule_precompute(project_id, mode=mode, is_user_action=True)

    return {"message": f"{mode} 模式预计算已进入调度队列"}


@router.post("/project/{project_id}/precompute/resume")
async def resume_precompute(
    project_id: str,
    body: PrecomputeRequest = PrecomputeRequest(),
    user: dict = Depends(get_current_user),
):
    """恢复预计算：不清除缓存，从断点续传。"""
    require_project_access(project_id, user, write=True)
    from core.precompute import schedule_precompute
    schedule_precompute(project_id, mode=body.mode, is_user_action=True)
    return {"message": f"{body.mode} 模式预计算已恢复调度"}


@router.post("/project/{project_id}/precompute/stop")
async def stop_precompute(
    project_id: str,
    body: PrecomputeRequest = PrecomputeRequest(),
    user: dict = Depends(get_current_user),
):
    """
    停止指定模式的运行中预计算任务。
    WHY: 管理员在发现预计算结果异常或需要重新配置时，
         需要能立即停止正在占用 GPU 的预计算任务。
    """
    require_project_access(project_id, user, write=True)
    mode = body.mode
    if mode not in ("generate", "replace", "clone"):
        raise HTTPException(400, f"无效模式: {mode}")

    from core.celery_app import celery_app
    from core.redis_client import get_redis

    r = get_redis()
    revoked = False
    if r:
        # 撤销 Celery 任务
        task_key = f"precompute:task_id:{project_id}:{mode}"
        task_id = r.get(task_key)
        if task_id:
            if isinstance(task_id, bytes):
                task_id = task_id.decode("utf-8")
            celery_app.control.revoke(task_id, terminate=True, signal="SIGTERM")
            revoked = True

        # 清除 Redis 运行状态
        r.delete(f"precompute:running:{project_id}:{mode}")
        r.delete(f"precompute:queued:{project_id}:{mode}")
        r.delete(task_key)
        r.delete(f"precompute:current_section:{project_id}:{mode}")

    logger.info(f"预计算已停止: project={project_id} mode={mode} revoked={revoked}")
    return {"message": f"{mode} 模式预计算已停止", "revoked": revoked}



# ---------- 预计算进度查询 ----------


@router.get("/project/{project_id}/precompute_status")
async def get_precompute_status(
    project_id: str, user: dict = Depends(get_current_user)
):
    """返回三个模式各自的预计算进度统计。"""
    require_project_access(project_id, user, write=False)
    from core.precompute import get_project_precompute_stats
    return get_project_precompute_stats(project_id)


# ---------- 草稿缓存读取 ----------


@router.get("/project/{project_id}/draft_cache/{mode}")
async def list_draft_cache(
    project_id: str, mode: str, user: dict = Depends(get_current_user)
):
    """
    批量读取指定模式下所有已缓存的章节草稿。
    WHY: 前端点击生成按钮时先调此接口判断是否有可用缓存。
    """
    require_project_access(project_id, user, write=False)
    if mode not in ("generate", "replace", "clone"):
        raise HTTPException(400, f"无效模式: {mode}")
    from core.precompute import get_all_draft_caches
    return get_all_draft_caches(project_id, mode)


@router.get("/project/{project_id}/draft_cache/{mode}/{section_id}")
async def get_draft_cache_item(
    project_id: str, mode: str, section_id: str,
    user: dict = Depends(get_current_user),
):
    """读取单章节缓存。"""
    require_project_access(project_id, user, write=False)
    from core.precompute import get_draft_cache
    cache = get_draft_cache(project_id, section_id, mode)
    if cache is None:
        raise HTTPException(404, "无缓存")
    return cache


@router.delete("/project/{project_id}/draft_cache")
async def clear_all_draft_cache(
    project_id: str, user: dict = Depends(get_current_user)
):
    """清除该项目所有模式的预计算缓存。"""
    require_project_access(project_id, user, write=True)
    from core.precompute import invalidate_draft_cache
    invalidate_draft_cache(project_id)
    return {"message": "全部预计算缓存已清除"}


@router.delete("/project/{project_id}/draft_cache/{mode}")
async def clear_mode_draft_cache(
    project_id: str, mode: str,
    user: dict = Depends(get_current_user),
):
    """清除指定模式的预计算缓存。"""
    require_project_access(project_id, user, write=True)
    if mode not in ("generate", "replace", "clone"):
        raise HTTPException(400, f"无效模式: {mode}")
    from core.precompute import invalidate_draft_cache
    invalidate_draft_cache(project_id, mode)
    return {"message": f"{mode} 模式缓存已清除"}


# ---------- 范文删除 ----------


@router.delete("/project/{project_id}")
async def delete_exemplar(
    project_id: str, user: dict = Depends(get_current_user)
):
    """删除指定项目绑定的写作范文，同时清理所有相关缓存。"""
    require_project_access(project_id, user, write=True)
    fp = EXEMPLARS_DIR / f"{project_id}.json"
    if fp.exists():
        fp.unlink()
        logger.info(f"范文已删除: project={project_id}")

    # WHY: 范文删除后 replace/clone 缓存彻底失效
    from core.precompute import invalidate_draft_cache
    invalidate_draft_cache(project_id)

    return {"message": "范文已删除"}

