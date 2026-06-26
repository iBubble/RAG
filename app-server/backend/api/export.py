import os
import re
import tempfile
from typing import List, Optional
from pydantic import BaseModel
from fastapi import APIRouter, HTTPException, Depends
from fastapi.responses import FileResponse, StreamingResponse
from celery.result import AsyncResult
import asyncio
from datetime import datetime, timezone, timedelta
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from core.auth_deps import get_current_user
from core.project_access import require_project_access
from core.docx_validator import validate_exported_docx
from core.docx_comments import inject_ai_comments
from core.docx_charts import parse_chart_marker, generate_chart_from_marker, CHART_MARKER_RE
from core.docx_cover import generate_cover_image
from core.docx_mermaid import render_mermaid_to_png

import logging
logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/export", tags=["出图与导出"])

class DocSection(BaseModel):
    title: str
    level: int  # 1, 2, 3 等，对应 一、 1. 1.1
    content: Optional[str] = None

class ExportRequest(BaseModel):
    project_id: str
    filename: str = "RAG_Generated_Report.docx"
    sections: List[DocSection]


# WHY: Markdown 表格行匹配正则（以 | 开头和结尾的行）
_MD_TABLE_LINE = re.compile(r'^\|(.+)\|$')
# WHY: 分隔符行匹配（如 | --- | --- |）
_MD_TABLE_SEP = re.compile(r'^\|\s*[-:]+\s*(\|\s*[-:]+\s*)*\|$')
# WHY: 图件占位标记匹配
_IMAGE_PLACEHOLDER = re.compile(r'\[插入图件[：:](.+?)\]')


def clean_collaborative_artifacts(content: str) -> str:
    """
    清除文档中在协同起草、质疑、大BOSS裁决过程中产生的过程性引导和思考词，
    确保成品文档仅包含最终的裁决正文（或流中断时的第一起草稿），且不破坏开头的大纲标题。
    """
    if not content:
        return content
    import re
    from api.admin import _read_system_settings
    sys_settings = _read_system_settings()
    contrarian_name = sys_settings.get("collab_contrarian_name", "【协同】审查员")
    arbiter_name = sys_settings.get("collab_arbiter_name", "【协同】仲裁官")
    esc_contrarian = re.escape(contrarian_name)
    esc_arbiter = re.escape(arbiter_name)
    
    expert_match = re.search(r'(?:⚖️|\[段落起草专家\])', content)
    boss_match = re.search(fr'(?:👑|\[大BOSS\]|\[{esc_arbiter}\]).*?(?:最终措辞润色|逻辑修正)', content)
    
    if expert_match and boss_match and boss_match.start() > expert_match.start():
        boss_end_pattern = fr'(?:👑|\[大BOSS\]|\[{esc_arbiter}\]).*?(?:最终措辞润色|逻辑修正).*?(?:</p>|\n)\s*'
        end_match = re.search(boss_end_pattern, content, flags=re.DOTALL)
        if end_match:
            header_part = content[:expert_match.start()].strip()
            header_part = re.sub(r'(?:⚖️|\u2696|\uFE0F|\s|\*|<strong>|<p>|<hr\s*/?>)+$', '', header_part).strip()
            body_part = content[end_match.end():].strip()
            
            # 如果大BOSS最终输出非空，则返回最终成品
            text_len = len(re.sub(r'<[^>]*>', '', body_part).strip())
            if text_len > 5:
                separator = '\n\n' if not header_part.endswith('>') else ''
                return (header_part + separator + body_part).strip()
                
    # 如果大BOSS最终输出为空，或者仅有初稿，则退回到只切除起草专家引导语并保留初稿
    if expert_match:
        expert_end_pattern = r'(?:⚖️|\[段落起草专家\]).*?正在起草章节初稿.*?(?:</p>|\n)\s*'
        end_match = re.search(expert_end_pattern, content, flags=re.DOTALL)
        if end_match:
            header_part = content[:expert_match.start()].strip()
            header_part = re.sub(r'(?:⚖️|\u2696|\uFE0F|\s|\*|<strong>|<p>|<hr\s*/?>)+$', '', header_part).strip()
            
            draft_part = content[end_match.end():].strip()
            # 如果有小杠或大BOSS引导语，切断之后的废话
            contrarian_match = re.search(fr'(?:🤨|\[小杠\]|\[{esc_contrarian}\]|👑|\[大BOSS\]|\[{esc_arbiter}\]|---|<blockquote|<hr)', draft_part)
            if contrarian_match:
                draft_part = draft_part[:contrarian_match.start()].strip()
                draft_part = re.sub(r'(?:⚖️|\u2696|\uFE0F|\s|\*|<strong>|<p>|<hr\s*/?>)+$', '', draft_part).strip()
                
            separator = '\n\n' if not header_part.endswith('>') else ''
            return (header_part + separator + draft_part).strip()
            
    return content


def _make_para_block(text: str) -> dict:
    """
    构建段落 block，自动检测 LaTeX 公式并拆分为 segments。
    WHY: 含 $...$ 的段落需要拆分为 text/omml 交替片段，
         C# 引擎据此生成 Word 原生数学公式。
    """
    if '$' in text:
        try:
            from core.latex_to_omml import split_text_and_math
            segments = split_text_and_math(text)
            if any(s['type'] == 'omml' for s in segments):
                return {'type': 'paragraph', 'segments': segments}
        except Exception as e:
            logger.warning(f'公式拆分失败，降级为纯文本: {e}')
    return {'type': 'paragraph', 'text': text}


def _set_cell_border(cell, **kwargs):
    """为单元格设置边框线。"""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('start', 'top', 'end', 'bottom'):
        if edge in kwargs:
            element = OxmlElement(f'w:{edge}')
            element.set(qn('w:val'), kwargs[edge].get('val', 'single'))
            element.set(qn('w:sz'), kwargs[edge].get('sz', '4'))
            element.set(qn('w:color'), kwargs[edge].get('color', '000000'))
            element.set(qn('w:space'), '0')
            tcBorders.append(element)
    tcPr.append(tcBorders)


def _add_cover_page(doc, report_title: str):
    """
    在文档开头插入封面页。
    WHY: 优先使用 Pillow 生成莫兰迪背景图封面，
         失败时降级为纯文本封面，不阻断导出。
    """
    cover_path = None
    try:
        cover_path = generate_cover_image(report_title)
        if cover_path:
            # WHY: 用全页图片作为封面，设置宽度为页面宽度
            section = doc.sections[0]
            page_width = section.page_width - section.left_margin - section.right_margin
            doc.add_picture(cover_path, width=page_width)
            # 分页
            page_break = doc.add_paragraph()
            run_br = page_break.add_run()
            run_br.add_break(WD_BREAK.PAGE)
            logger.info("莫兰迪封面已插入")
            return
    except Exception as exc:
        logger.warning(f"封面图生成/插入异常，降级为纯文本封面: {exc}")
    finally:
        # WHY: 清理临时文件
        if cover_path:
            try:
                os.remove(cover_path)
            except OSError:
                pass

    # ── 降级：纯文本封面 ──
    for _ in range(6):
        doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(6)
    title_run = title_p.add_run(report_title)
    title_run.font.name = 'SimHei'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
    title_run.font.size = Pt(26)
    title_run.bold = True

    line_p = doc.add_paragraph()
    line_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line_p.paragraph_format.space_before = Pt(12)
    line_p.paragraph_format.space_after = Pt(12)
    line_run = line_p.add_run('\u2501' * 30)
    line_run.font.size = Pt(10)
    line_run.font.color.rgb = RGBColor(0x33, 0x66, 0x99)

    for _ in range(4):
        doc.add_paragraph()

    org_p = doc.add_paragraph()
    org_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    org_run = org_p.add_run('智能体')
    org_run.font.name = 'FangSong'
    org_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'FangSong')
    org_run.font.size = Pt(18)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.paragraph_format.space_before = Pt(8)
    date_str = datetime.now(timezone(timedelta(hours=8))).strftime('%Y年%m月')
    date_run = date_p.add_run(date_str)
    date_run.font.name = 'FangSong'
    date_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'FangSong')
    date_run.font.size = Pt(16)

    page_break = doc.add_paragraph()
    run_br = page_break.add_run()
    run_br.add_break(WD_BREAK.PAGE)


def _add_header_footer(doc, header_text: str):
    """
    为文档添加页眉和页脚。
    WHY: 专业报告需要页眉标题 + 页脚页码。
    """
    section = doc.sections[0]

    header = section.header
    header.is_linked_to_previous = False
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pPr = header_p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom_border = OxmlElement('w:bottom')
    bottom_border.set(qn('w:val'), 'single')
    bottom_border.set(qn('w:sz'), '6')
    bottom_border.set(qn('w:space'), '1')
    bottom_border.set(qn('w:color'), '336699')
    pBdr.append(bottom_border)
    pPr.append(pBdr)
    run = header_p.add_run(header_text)
    run.font.name = 'FangSong'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'FangSong')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    footer = section.footer
    footer.is_linked_to_previous = False
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _footer_run(text):
        r = footer_p.add_run(text)
        r.font.name = 'FangSong'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'FangSong')
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    _footer_run('\u7b2c ')
    fld_page = OxmlElement('w:fldSimple')
    fld_page.set(qn('w:instr'), 'PAGE')
    pr = OxmlElement('w:r')
    pt_ = OxmlElement('w:t')
    pt_.text = '1'
    pr.append(pt_)
    fld_page.append(pr)
    footer_p._element.append(fld_page)
    _footer_run(' \u9875  \u5171 ')
    fld_total = OxmlElement('w:fldSimple')
    fld_total.set(qn('w:instr'), 'NUMPAGES')
    tr = OxmlElement('w:r')
    tt = OxmlElement('w:t')
    tt.text = '1'
    tr.append(tt)
    fld_total.append(tr)
    footer_p._element.append(fld_total)
    _footer_run(' \u9875')


def _add_markdown_table(doc, md_lines: List[str]):
    """
    将 Markdown 表格语法转为 Word 专业三线表。
    WHY: 三线表（顶线、底线加粗，表头下细线，无竖线）是
         土地评估/工程咨询行业的标准表格样式。
         ≥6 列的宽表自动切换为横向页面排版。
    """
    from docx.enum.section import WD_ORIENT

    data_lines = [
        line for line in md_lines
        if not _MD_TABLE_SEP.match(line.strip())
    ]
    if not data_lines:
        return

    rows = []
    for line in data_lines:
        cells = [c.strip() for c in line.strip('|').split('|')]
        rows.append(cells)

    if not rows:
        return

    n_cols = max(len(r) for r in rows)
    n_rows = len(rows)

    # WHY: ≥6 列的宽表在 A4 纵向放不下，切换为横向排版
    need_landscape = n_cols >= 6

    if need_landscape:
        # 插入分节符，切换为横向页面
        new_section = doc.add_section(2)  # 2 = WD_SECTION_START.NEW_PAGE
        new_section.orientation = WD_ORIENT.LANDSCAPE
        new_section.page_width = Cm(29.7)   # A4 高度变宽度
        new_section.page_height = Cm(21.0)  # A4 宽度变高度
        new_section.left_margin = Cm(2.54)
        new_section.right_margin = Cm(2.54)
        new_section.top_margin = Cm(1.91)
        new_section.bottom_margin = Cm(1.91)

    # 获取当前页面可用宽度
    current_section = doc.sections[-1]
    page_avail = current_section.page_width - current_section.left_margin - current_section.right_margin

    # 根据每列最大内容长度计算列宽权重
    col_weights = []
    for j in range(n_cols):
        max_w = 0
        for row_data in rows:
            if j < len(row_data):
                text = row_data[j]
                w = sum(2 if ord(c) > 127 else 1 for c in text)
                max_w = max(max_w, w)
        col_weights.append(max(4, min(max_w, 60)))

    total_weight = sum(col_weights)

    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = 'Table Grid'
    table.autofit = False

    # 设置各列宽度
    for j in range(n_cols):
        col_width = int(page_avail * col_weights[j] / total_weight)
        for i in range(n_rows):
            table.cell(i, j).width = col_width

    # 列数较多时缩小字号
    if n_cols >= 10:
        cell_font_size = Pt(8)
    elif n_cols >= 6:
        cell_font_size = Pt(9)
    else:
        cell_font_size = Pt(12)

    THICK = {'val': 'single', 'sz': '12', 'color': '000000'}
    THIN = {'val': 'single', 'sz': '6', 'color': '000000'}
    NONE = {'val': 'none', 'sz': '0', 'color': 'auto'}

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j < n_cols:
                cell = table.cell(i, j)
                cell.text = cell_text
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.font.size = cell_font_size
                        run.font.name = 'FangSong'
                        run._element.rPr.rFonts.set(
                            qn('w:eastAsia'), 'FangSong'
                        )
                    if i == 0:
                        for run in p.runs:
                            run.bold = True

                borders = {'start': NONE, 'end': NONE}
                if i == 0:
                    borders['top'] = THICK
                    borders['bottom'] = THIN
                elif i == n_rows - 1:
                    borders['top'] = NONE
                    borders['bottom'] = THICK
                else:
                    borders['top'] = NONE
                    borders['bottom'] = NONE
                _set_cell_border(cell, **borders)

    if need_landscape:
        # 表格结束后，切回纵向页面
        new_section = doc.add_section(2)
        new_section.orientation = WD_ORIENT.PORTRAIT
        new_section.page_width = Cm(21.0)
        new_section.page_height = Cm(29.7)
        new_section.left_margin = Cm(3.17)
        new_section.right_margin = Cm(3.17)
        new_section.top_margin = Cm(2.54)
        new_section.bottom_margin = Cm(2.54)
    else:
        doc.add_paragraph()


def _add_image_placeholder(doc, placeholder_text: str):
    """
    将 [插入图件：...] 占位标记转为黄色高亮段落。
    WHY: 提示用户在此处手工插入图件。
    """
    p = doc.add_paragraph()
    run = p.add_run(f"⚠️ 此处需手工插入图件 — {placeholder_text}")
    run.font.size = Pt(12)
    run.font.name = 'FangSong'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'FangSong')
    run.font.color.rgb = RGBColor(0x8B, 0x6B, 0x00)
    # 设置黄色底纹
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'FFFFCC')
    run._element.rPr.append(shd)


def _render_section_content(doc, content: str):
    """
    解析章节正文内容，分别处理：
    - 可视化标记 → 解析表格数据 → 生成图表图片 → 插入 Word
    - 普通段落 → Word 段落（仿宋四号，首行缩进）
    - Markdown 表格 → Word 专业三线表
    - 图件占位标记 → 黄底高亮提示
    """
    if not content or not content.strip():
        return

    # WHY: 在逐行处理之前，先用正则一次性剥离所有 ```mermaid...``` 代码块。
    #       某些 Mermaid 代码块是内联拼在段落里的（不是单独一行），
    #       逐行匹配 "```mermaid" 会漏掉这种情况。正则兜底最可靠。
    content = re.sub(r'```mermaid\b.*?```', '', content, flags=re.DOTALL)

    lines = content.strip().split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()

        # WHY: 检测 [可视化：柱状图，标题：XXX] 标记
        chart_info = parse_chart_marker(line)
        if chart_info:
            i += 1
            # 收集紧随标记的 Markdown 表格行
            table_lines = []
            while i < len(lines) and _MD_TABLE_LINE.match(lines[i].strip()):
                table_lines.append(lines[i].strip())
                i += 1
            if table_lines:
                chart_path = generate_chart_from_marker(chart_info, table_lines)
                if chart_path:
                    try:
                        # 插入图表图片（宽度 14cm，约占页面 80%）
                        doc.add_picture(chart_path, width=Cm(14))
                        # 图表标题（居中，小五号）
                        cap = doc.add_paragraph()
                        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        cap_run = cap.add_run(f"图  {chart_info['title']}")
                        cap_run.font.size = Pt(9)
                        cap_run.font.name = 'SimSun'
                        cap_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
                        cap_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                        logger.info(f"图表已插入: {chart_info['title']}")
                    except Exception as exc:
                        logger.error(f"图表插入失败: {exc}")
                    finally:
                        # WHY: 清理临时文件
                        try:
                            os.remove(chart_path)
                        except OSError:
                            pass
                # 同时也输出原始表格（数据留底）
                _add_markdown_table(doc, table_lines)
            continue

        # 检测 Markdown 表格块（连续多行以 | 开头的）
        if _MD_TABLE_LINE.match(line):
            table_lines = []
            while i < len(lines) and _MD_TABLE_LINE.match(lines[i].strip()):
                table_lines.append(lines[i].strip())
                i += 1
            _add_markdown_table(doc, table_lines)
            continue

        # 检测 Mermaid 代码块（```mermaid ... ```）
        # WHY: Word 无法渲染 Mermaid，之前的"降级"会把原始代码塞进文档变成乱码。
        #       现在直接整块跳过，静默丢弃。
        if line == '```mermaid':
            i += 1
            while i < len(lines) and lines[i].strip() != '```':
                i += 1
            if i < len(lines):
                i += 1  # 跳过闭合 ```
            continue

        # 检测图件占位标记
        img_match = _IMAGE_PLACEHOLDER.search(line)
        if img_match:
            _add_image_placeholder(doc, img_match.group(0))
            i += 1
            continue

        # 普通段落（跳过空行）
        if line:
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Pt(28)
            run = p.add_run(line)
            run.font.size = Pt(14)
            run.font.name = 'FangSong'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'FangSong')

        i += 1


@router.post("/docx")
async def export_to_docx(req: ExportRequest, user: dict = Depends(get_current_user)):
    """
    根据前端传来的文档骨架与生成的段落内容，
    使用标准工程公文排版（黑体/仿宋），组装成真实的 .docx 文件并返回下载。
    增强功能：Markdown 表格 → Word 表格，图件占位 → 黄底高亮段落。
    导出后自动执行质检（表格列宽、图片变形、空章节、字体一致性）。
    WHY: 导出是读操作，公开项目允许所有登录用户使用。
    """
    require_project_access(req.project_id, user, write=False)
    if not req.sections:
        raise HTTPException(status_code=400, detail="文档内容为空")

    # WHY: 记录导出操作到审计日志
    from core.audit_log import log_operation
    log_operation(user["id"], "document_export", f"导出文档：{req.filename}（{len(req.sections)}个章节）")

    import tempfile
    import os
    import subprocess
    import json
    import html as html_lib
    import uuid

    report_title = req.filename.replace('.docx', '').replace('.DOCX', '')

    # 1. 构建适合 C# SDK 的 JSON 模型
    doc_request = {
        "title": report_title,
        "org_name": "智能体",
        "date_str": datetime.now(timezone(timedelta(hours=8))).strftime("%Y年%m月"),
        "sections": []
    }
    for sec in req.sections:
        clean_content = ""
        if sec.content and sec.content.strip():
            sec.content = clean_collaborative_artifacts(sec.content)
            logger.info(f"[EXPORT_DEBUG] SECTION TITLE: {sec.title}, content_len={len(sec.content)}, has_katex={'katex' in sec.content.lower()}")
            logger.info(f"[EXPORT_DEBUG] RAW CONTENT SAMPLE: {sec.content[:300]}")
            
            # WHY: 前端传来的富文本通常包含 <table> 标签。由于后续正则不仅会去掉 table 和 tr，
            #      还会把 td 直接抹掉而没有空格符，导致单元格粘连失去列结构。
            #      所以我们在正则剥离 HTML 前，先用 BS4 把 table 标签手动转成 \n| A | B |\n 的 Markdown 表格结构。
            try:
                from bs4 import BeautifulSoup
                soup = BeautifulSoup(sec.content, 'html.parser')

                # WHY: 前端编辑器（TipTap）将 KaTeX 渲染结果包裹在自定义容器中：
                #   块级: <div class="katex-block">&lt;span class="katex-display"&gt;...&lt;/span&gt;</div>
                #   行内: <span class="katex-inline">&lt;span class="katex"&gt;...&lt;/span&gt;</span>
                # 注意：内部 KaTeX HTML 被 HTML 实体转义了（&lt; &gt; &amp;），
                #       BS4 看到的是纯文本而非 HTML 标签。
                # 策略：找到外层容器 → unescape 内容 → 重新解析 → 提取 annotation 中的 LaTeX 源码。
                import html as html_module

                # 处理块级公式 <div class="katex-block">
                for container in soup.find_all(['div', 'span'], class_='katex-block'):
                    inner_escaped = container.decode_contents()
                    inner_html = html_module.unescape(inner_escaped)
                    inner_soup = BeautifulSoup(inner_html, 'html.parser')
                    annotation = inner_soup.find(
                        'annotation', attrs={'encoding': 'application/x-tex'}
                    )
                    if annotation and annotation.string and annotation.string.strip():
                        latex_src = annotation.string.strip()
                        container.replace_with(f' $${latex_src}$$ ')
                        logger.info(f"[KATEX] 块级公式还原: {latex_src[:60]}")
                    else:
                        container.decompose()

                # 处理行内公式 <span class="katex-inline">
                for container in soup.find_all('span', class_='katex-inline'):
                    inner_escaped = container.decode_contents()
                    inner_html = html_module.unescape(inner_escaped)
                    inner_soup = BeautifulSoup(inner_html, 'html.parser')
                    annotation = inner_soup.find(
                        'annotation', attrs={'encoding': 'application/x-tex'}
                    )
                    if annotation and annotation.string and annotation.string.strip():
                        latex_src = annotation.string.strip()
                        container.replace_with(f' ${latex_src}$ ')
                        logger.info(f"[KATEX] 行内公式还原: {latex_src[:60]}")
                    else:
                        container.decompose()

                # 兜底：处理未被包裹的原生 KaTeX span（直接渲染的场景）
                for katex_block in soup.find_all('span', class_='katex-display'):
                    annotation = katex_block.find(
                        'annotation', attrs={'encoding': 'application/x-tex'}
                    )
                    if annotation and annotation.string:
                        katex_block.replace_with(f' $${annotation.string.strip()}$$ ')
                    else:
                        katex_block.decompose()

                for katex_inline in soup.find_all('span', class_='katex'):
                    annotation = katex_inline.find(
                        'annotation', attrs={'encoding': 'application/x-tex'}
                    )
                    if annotation and annotation.string:
                        katex_inline.replace_with(f' ${annotation.string.strip()}$ ')
                    else:
                        katex_inline.decompose()

                for table in soup.find_all('table'):
                    md_table = []
                    for row in table.find_all('tr'):
                        cols = row.find_all(['td', 'th'])
                        # 提取文本，将单元格内的换行替换为空格，避免打断 Markdown 表格的单行结构
                        row_text = "| " + " | ".join(
                            c.get_text(separator=" ", strip=True).replace("\n", " ").replace("|", "｜") for c in cols
                        ) + " |"
                        md_table.append(row_text)
                    if md_table:
                        table.replace_with("\n\n" + "\n".join(md_table) + "\n\n")
                # 重新赋回 content，以便后续复用正则清理逻辑
                sec.content = str(soup)
            except ImportError:
                pass

            clean_content = re.sub(r'(?i)<(?:br|/?(?:p|div|li|ul|ol|h[1-6]|table|tbody|thead|tr))[^>]*>', '\n', sec.content)
            clean_content = re.sub(r'<[^>]*>?', '', clean_content)
            clean_content = html_lib.unescape(clean_content)
            # WHY: 在 HTML 剥离后，统一清除所有残留的 ```mermaid...``` 代码块。
            #       旧内容里的 Mermaid 可能内嵌在 <p> 标签里，到此处已变成纯文本。
            clean_content = re.sub(r'```mermaid\b.*?```', '', clean_content, flags=re.DOTALL)
            clean_content = re.sub(r'(\[可视化[：:].+?\])', r'\n\1\n', clean_content)
            # WHY: 原先此处有 re.sub(r'(\n|^)(\|.*\|)(?=\n|$)', r'\n\2\n', ...)
            #       它给每个 | 开头的表格行前后都加了空行，导致连续的 Markdown 表格行
            #       被空行打断，解析器把每一行当成独立表格。已移除。
            
            # WHY: 去掉内容开头与章节标题重复的文本行
            # 前端传来的 content 经常以标题本身开头（如 "一、课程概述\n正文..."），
            # 而 C# 引擎已经会用 sec.title 生成独立的 Heading，不去掉就会重复。
            content_lines = clean_content.split('\n')
            stripped_lines = []
            title_stripped = False
            for cl in content_lines:
                cl_clean = cl.strip()
                if not title_stripped and cl_clean:
                    # 如果这一行和标题一模一样，或者去掉序号后一样，就跳过
                    normalized_title = re.sub(r'^[一二三四五六七八九十\d（\(）\)、\.\s]+', '', sec.title).strip()
                    normalized_line = re.sub(r'^[一二三四五六七八九十\d（\(）\)、\.\s]+', '', cl_clean).strip()
                    if normalized_line == normalized_title or cl_clean == sec.title:
                        title_stripped = True
                        continue
                if cl_clean or title_stripped:
                    title_stripped = True  # 一旦跳过了标题行或遇到非空行，标记完成
                    stripped_lines.append(cl)
            clean_content = '\n'.join(stripped_lines)
        
        blocks = []
        if clean_content:
            # DEBUG: 输出 clean_content 中的表格行，排查表格被拆分的原因
            table_debug_lines = [l for l in clean_content.split('\n') if '|' in l]
            if table_debug_lines:
                logger.info(f"[TABLE_DEBUG] section={sec.title}, table_lines_count={len(table_debug_lines)}")
                for idx, tdl in enumerate(table_debug_lines[:5]):
                    logger.info(f"[TABLE_DEBUG]   line[{idx}]: {tdl[:120]}")
                # 输出表格行周围的上下文——确认是否有空行打断连续性
                all_lines = clean_content.split('\n')
                for idx, al in enumerate(all_lines):
                    if '|' in al and idx > 0:
                        prev = all_lines[idx-1]
                        logger.info(f"[TABLE_DEBUG]   prev_of_pipe[{idx}]: '{prev[:80]}' (empty={prev.strip()==''})") 
                        break

            # WHY: 把表格行之间的空行去掉，确保 |...| 行连续，不被解析器拆成多个独立表格。
            #       同时合并连续多个空行为一个，减少噪音。
            clean_lines = clean_content.split('\n')
            merged = []
            for ci, cl in enumerate(clean_lines):
                if cl.strip() == '':
                    # 如果前一行和后一行都是 | 开头的表格行，跳过这个空行
                    prev_is_pipe = merged and merged[-1].strip().startswith('|')
                    # 往后找下一个非空行
                    next_is_pipe = False
                    for nj in range(ci + 1, len(clean_lines)):
                        if clean_lines[nj].strip():
                            next_is_pipe = clean_lines[nj].strip().startswith('|')
                            break
                    if prev_is_pipe and next_is_pipe:
                        continue  # 跳过表格行之间的空行
                merged.append(cl)
            clean_content = '\n'.join(merged)

            lines = clean_content.split('\n')

            # ── 管道符表格预清洗 ──
            # WHY: LLM 生成的概算表等复杂表格常有格式缺陷：
            #   - 双管道 `||` → 空单元格导致连续管道符
            #   - 分类行只有一列有值，其余为空 `| 第一部分 建筑工程 ||||| |`
            #   - 列数不一致（某些行少列或多列）
            #   - 管道符行被文本换行拆断（如 `| 独立\n费用 | 合计 |`）
            #   这些异常会导致解析失败，表格退化为纯文本。

            # 步骤0：断行拼接 — 将被换行拆断的管道符行重新合并
            # WHY: 编辑器中的文本可能因自动换行导致一个 `|...|` 行被拆成两行。
            #       特征：当前行以 `|` 开头但下一行不以 `|` 开头（且非空），说明被截断了。
            joined_lines = []
            i = 0
            while i < len(lines):
                s = lines[i].strip()
                if s.startswith('|') and not s.endswith('|'):
                    # 向后找续行，直到遇到以 `|` 结尾的行或以 `|` 开头的新行
                    combined = s
                    while i + 1 < len(lines):
                        next_s = lines[i + 1].strip()
                        if next_s.startswith('|'):
                            break  # 下一行是新的管道符行
                        if not next_s:
                            break  # 空行，停止拼接
                        combined += ' ' + next_s
                        i += 1
                        if combined.endswith('|'):
                            break
                    joined_lines.append(combined)
                else:
                    joined_lines.append(lines[i])
                i += 1
            lines = joined_lines

            # 步骤1-2：双管道标准化 + 确保行尾有 `|`
            sanitized_lines = []
            for line in lines:
                s = line.strip()
                if s.startswith('|'):
                    while '||' in s:
                        s = s.replace('||', '| |')
                    if not s.endswith('|'):
                        s += ' |'
                    sanitized_lines.append(s)
                else:
                    sanitized_lines.append(line)
            lines = sanitized_lines

            current_p = []
            in_table = False
            table_rows = []
            header_col_count = 0  # 表头列数，用于对齐后续行
            
            for line in lines:
                striped = line.strip()
                if striped.startswith("|"):
                    if not in_table:
                        if current_p:
                            blocks.append(_make_para_block(" ".join(current_p).strip()))
                            current_p = []
                        in_table = True
                        header_col_count = 0  # 新表格重置
                    row = [col.strip() for col in striped.strip("|").split("|")]

                    # 分隔行检测（跳过）：容忍空单元格和额外空格
                    non_empty_cells = [r for r in row if r.strip()]
                    if (len(row) > 0
                            and all(all(c in "- : " for c in r) for r in row)
                            and striped.count("-") >= 3):
                        continue

                    if row:
                        # 记录表头列数（第一个数据行）
                        if header_col_count == 0:
                            header_col_count = len(row)

                        # 步骤3：列数对齐 — padding 不足的行，截断多余的列
                        if header_col_count > 0:
                            if len(row) < header_col_count:
                                row.extend([''] * (header_col_count - len(row)))
                            elif len(row) > header_col_count:
                                row = row[:header_col_count]

                        # 步骤4：跳过全空行（所有单元格都为空）
                        if all(cell.strip() == '' for cell in row):
                            continue

                        table_rows.append(row)
                elif striped.startswith("[可视化：") or striped.startswith("[可视化:"):
                    if current_p:
                        blocks.append(_make_para_block(" ".join(current_p).strip()))
                        current_p = []
                    chart_title = "数据图表"
                    m = re.search(r'可视化[：:](.*?)\]', striped)
                    if m: chart_title = m.group(1).strip()
                    c_type = "pie" if "占比" in chart_title or "分布" in chart_title else "bar"
                    blocks.append({
                        "type": "chart", 
                        "chartType": c_type,
                        "chartTitle": chart_title,
                        "chartData": []
                    })
                else:
                    if in_table:
                        in_table = False
                        # 步骤5：空列折叠 — 移除所有行中都为空的列
                        if table_rows and len(table_rows) > 1:
                            col_count = len(table_rows[0])
                            non_empty_cols = []
                            for ci in range(col_count):
                                # 如果该列在所有行（含表头）中都为空 → 折叠
                                if any(ci < len(r) and r[ci].strip() for r in table_rows):
                                    non_empty_cols.append(ci)
                            if len(non_empty_cols) < col_count:
                                table_rows = [
                                    [r[ci] if ci < len(r) else '' for ci in non_empty_cols]
                                    for r in table_rows
                                ]
                        blocks.append({"type": "table", "rows": table_rows})
                        table_rows = []
                        header_col_count = 0
                    if striped == "":
                        if current_p:
                            blocks.append(_make_para_block(" ".join(current_p).strip()))
                            current_p = []
                    else:
                        current_p.append(striped)
            if current_p:
                blocks.append(_make_para_block(" ".join(current_p).strip()))
            if in_table and table_rows:
                # 末尾表格也做空列折叠
                if len(table_rows) > 1:
                    col_count = len(table_rows[0])
                    non_empty_cols = []
                    for ci in range(col_count):
                        if any(ci < len(r) and r[ci].strip() for r in table_rows):
                            non_empty_cols.append(ci)
                    if len(non_empty_cols) < col_count:
                        table_rows = [
                            [r[ci] if ci < len(r) else '' for ci in non_empty_cols]
                            for r in table_rows
                        ]
                blocks.append({"type": "table", "rows": table_rows})
        
        doc_request["sections"].append({
            "title": sec.title,
            "level": sec.level,
            "blocks": blocks
        })

    # 2. 写入临时 JSON 并调用 C# 引擎
    tmp_dir = tempfile.gettempdir()
    run_id = str(uuid.uuid4())
    json_path = os.path.join(tmp_dir, f"export_{run_id}.json")
    tmp_path = os.path.join(tmp_dir, f"export_{run_id}.docx")
    
    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump(doc_request, f, ensure_ascii=False)
        
    try:
        from worker import generate_docx_bg
        # Fire and forget asynchronously
        task = generate_docx_bg.delay(run_id, json_path, tmp_path)
        return {"task_id": task.id, "message": "Docx Generation Started"}
        
    except Exception as e:
        if os.path.exists(json_path):
            os.remove(json_path)
        raise HTTPException(status_code=500, detail=f"提交后台引擎失败: {str(e)}")

@router.get("/docx-status")
async def docx_status(task_id: str, user: dict = Depends(get_current_user)):
    """
    轮询式进度查询接口（替代 SSE 流）。
    WHY: FRP HTTP 代理会缓冲整个 SSE 响应，导致浏览器 EventSource 收不到实时事件。
         改为普通 JSON 轮询，100% 兼容所有代理层（FRP / Nginx / Caddy）。
    """
    res = AsyncResult(task_id)
    if res.ready():
        if res.successful():
            return {"percent": 100, "message": "完毕", "status": "success"}
        else:
            return {"percent": -1, "message": "错误发生", "status": "failed"}
    
    if res.state == "PROGRESS":
        info = res.info or {}
        return {
            "percent": info.get("percent", 0),
            "message": info.get("message", "生成中..."),
            "status": "progress"
        }
    
    return {"percent": 0, "message": "排队中...", "status": "pending"}

@router.get("/docx-progress")
async def docx_progress(task_id: str, user: dict = Depends(get_current_user)):
    """SSE流式接口：透传后台生成进度（保留兼容，推荐使用 /docx-status 轮询）"""
    async def event_generator():
        last_percent = -1
        while True:
            res = AsyncResult(task_id)
            if res.ready():
                if res.successful():
                    yield 'data: {"percent": 100, "message": "完毕", "status": "success"}\n\n'
                else:
                    yield 'data: {"percent": -1, "message": "错误发生", "status": "failed"}\n\n'
                await asyncio.sleep(0.5)
                break

            if res.state == "PROGRESS":
                info = res.info or {}
                pct = info.get("percent", 0)
                msg = info.get("message", "生成中...")
                if pct != last_percent:
                    last_percent = pct
                    import json
                    yield f"data: {json.dumps({'percent': pct, 'message': msg, 'status': 'progress'})}\n\n"
            
            await asyncio.sleep(0.5)

    return StreamingResponse(event_generator(), media_type="text/event-stream", headers={
        "Cache-Control": "no-cache",
        "Connection": "keep-alive"
    })

@router.get("/download/{task_id}")
async def download_docx(task_id: str, user: dict = Depends(get_current_user)):
    res = AsyncResult(task_id)
    if not res.successful():
        raise HTTPException(400, "文件尚未生成或者生成失败")
    
    out_dict = res.result
    tmp_path = out_dict.get("file_path")
    if not tmp_path or not os.path.exists(tmp_path):
        raise HTTPException(404, "文件不存在")

    # TODO QC header isn't easily passed over task Result, let's keep it simple for now
    return FileResponse(
        path=tmp_path,
        filename="Exported_Document.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


class CoverPreviewRequest(BaseModel):
    title: str
    org_name: Optional[str] = "智能体"
    date_str: Optional[str] = None

@router.post("/cover-preview")
async def cover_preview(req: CoverPreviewRequest, user: dict = Depends(get_current_user)):
    """
    生成封面预览缩略图（低分辨率），返回 base64 编码的 PNG。
    WHY: 前端导出确认弹窗需要展示封面效果预览，
         让用户在下载前直观确认文档外观。
    """
    import base64
    from PIL import Image as PILImage

    cover_path = None
    try:
        cover_path = generate_cover_image(req.title, req.org_name, req.date_str)
        if not cover_path:
            raise HTTPException(status_code=500, detail="封面生成失败")

        # WHY: 缩小到宽 400px 的缩略图，减少传输体积
        with PILImage.open(cover_path) as img:
            ratio = 400 / img.width
            thumb_size = (400, int(img.height * ratio))
            img_thumb = img.resize(thumb_size, PILImage.LANCZOS)

            import io
            buf = io.BytesIO()
            img_thumb.save(buf, format="PNG", optimize=True)
            b64 = base64.b64encode(buf.getvalue()).decode("ascii")

        return {"image": f"data:image/png;base64,{b64}"}
    except HTTPException:
        raise
    except Exception as exc:
        logger.error(f"封面预览生成失败: {exc}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"封面预览生成失败: {exc}")
    finally:
        if cover_path:
            try:
                os.remove(cover_path)
            except OSError:
                pass


class InternalAnnotateRequest(BaseModel):
    file_path: str

@router.post("/internal/docx/annotate")
async def internal_docx_annotate(req: InternalAnnotateRequest):
    """
    内部 Docx 留痕批注注入微服务接口。
    """
    from core.docx_comments import inject_ai_comments
    if not os.path.exists(req.file_path):
        raise HTTPException(status_code=400, detail="指定文件不存在")
    try:
        inject_ai_comments(req.file_path)
        return {"status": "success"}
    except Exception as e:
        logger.error(f"注入批注留痕失败: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


