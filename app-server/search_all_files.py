import os
from pathlib import Path
import docx

def search_docx(file_path, term):
    try:
        doc = docx.Document(file_path)
        for i, p in enumerate(doc.paragraphs):
            if term in p.text:
                print(f"Found in DOCX paragraph: {file_path.name}, p {i}: {p.text}")
        for t_idx, table in enumerate(doc.tables):
            for r_idx, row in enumerate(table.rows):
                row_text = " | ".join(cell.text for cell in row.cells)
                if term in row_text:
                    print(f"Found in DOCX table: {file_path.name}, table {t_idx}, row {r_idx}: {row_text}")
    except Exception as e:
        pass

def search_text_or_json(file_path, term):
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        if term in content:
            print(f"Found in Text/JSON: {file_path}")
            # Print matching lines
            lines = content.split('\n')
            for i, line in enumerate(lines):
                if term in line:
                    print(f"  Line {i}: {line[:150]}")
    except Exception as e:
        pass

term = "3323.63"
print(f"Searching for '{term}'...")

root_dir = Path("/Volumes/SYRAID/RAG_Files")
for p in root_dir.rglob("*"):
    if p.is_file():
        ext = p.suffix.lower()
        if ext == ".docx":
            search_docx(p, term)
        elif ext in (".xlsx", ".xls"):
            # Already searched, but let's double check if we missed any
            pass
        elif ext in (".json", ".txt", ".md", ".csv"):
            search_text_or_json(p, term)

print("Search completed.")
